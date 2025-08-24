"""
Automação Avançada para Geração de Documentação de Projetos Power BI
---------------------------------------------------------------------

Este código automatiza a documentação de relatórios Power BI a partir de um arquivo .pbit.
Ele extrai informações detalhadas e utiliza a API do Google Gemini para gerar análises
inteligentes para cada página e tabela do relatório, enriquecendo a documentação final
gerada em um arquivo Word.

Versão: 2.0 (Com integração de IA por item)
"""
from dotenv import load_dotenv
import json
import os
import zipfile
from datetime import datetime
from docx import Document
import google.generativeai as genai
import config as cfg


# ===================================================================
# FUNÇÕES UTILITÁRIAS
# ===================================================================

def verificar_ou_renomear_arquivo(arquivo_pbit, arquivo_zip):
    """Verifica se o arquivo .zip já existe; se não, renomeia o .pbit."""
    if os.path.exists(arquivo_zip):
        print("Arquivo .zip já existe. Pulando renomeação.")
    else:
        print(f"Renomeando '{arquivo_pbit}' para '{arquivo_zip}'.")
        os.rename(arquivo_pbit, arquivo_zip)


def extrair_arquivos_zip(arquivo_zip, caminho_bi, arquivos_para_extrair):
    """Extrai arquivos específicos de um arquivo ZIP."""
    print(f"Extraindo arquivos de '{arquivo_zip}'...")
    with zipfile.ZipFile(arquivo_zip, 'r') as zip_ref:
        for arquivo in arquivos_para_extrair:
            zip_ref.extract(arquivo, caminho_bi)
    print("Arquivos extraídos com sucesso.")


def carregar_dados_json(arquivo: str, encoding: str = 'utf-16-le') -> dict:
    """Carrega dados de um arquivo JSON."""
    try:
        with open(arquivo, 'r', encoding=encoding) as f:
            return json.load(f)
    except Exception as e:
        print(f"Erro ao carregar JSON: {arquivo} - {e}")
        return {}


# ===================================================================
# FUNÇÕES DE EXTRAÇÃO DE DADOS ESTRUTURADOS
# ===================================================================

def extrair_dados_paginas(layout: dict) -> list:
    """Extrai dados de páginas e seus visuais, retornando uma lista de páginas."""
    paginas = []
    for section in layout.get('sections', []):
        pagina_info = {
            "nome": section.get('displayName', 'Sem Nome'),
            "visuais": []
        }
        for container in section.get("visualContainers", []):
            try:
                config_data = json.loads(container.get("config", "{}"))
                visual_type = config_data.get("singleVisual", {}).get("visualType")
                medidas_usadas = [item.get("queryRef") for items in
                                  config_data.get("singleVisual", {}).get("projections", {}).values()
                                  for item in items if item.get("queryRef")]
                pagina_info["visuais"].append({
                    "tipo": visual_type,
                    "medidas": medidas_usadas if medidas_usadas else ["Nenhuma medida explícita"]
                })
            except (json.JSONDecodeError, AttributeError):
                continue  # Ignora visuais com config malformado
        paginas.append(pagina_info)
    return paginas


def extrair_dados_tabelas(model_data: dict) -> list:
    """Extrai dados de tabelas, colunas e medidas, retornando uma lista de tabelas."""
    tabelas = []
    for table in model_data.get('model', {}).get('tables', []):
        if table.get("name", "").startswith(("DateTableTemplate", "LocalDateTable")):
            continue
        tabela_info = {
            "nome": table.get("name", ""),
            "colunas": [],
            "medidas": []
        }
        for column in table.get('columns', []):
            tabela_info["colunas"].append({
                "nome": column.get("name", ""),
                "tipo_dado": column.get('dataType', ""),
                "calculada": 'Sim' if column.get('type', "") in ['calculatedTableColumn', 'calculated'] else 'Não'
            })
        for measure in table.get('measures', []):
            expression = measure.get('expression', '')
            if isinstance(expression, list):
                expression = ' '.join(filter(lambda x: x.strip(), expression))
            tabela_info["medidas"].append({
                "nome": measure.get('name', ''),
                "expressao": expression
            })
        tabelas.append(tabela_info)
    return tabelas


def extrair_dados_relacionamentos(model_data: dict) -> list:
    """Extrai dados de relacionamentos, retornando uma lista."""
    relacionamentos = []
    for relation in model_data.get('model', {}).get('relationships', []):
        from_table = relation.get('fromTable')
        to_table = relation.get('toTable')
        if from_table.startswith(("DateTableTemplate", "LocalDateTable")) or \
                to_table.startswith(("DateTableTemplate", "LocalDateTable")):
            continue
        relacionamentos.append({
            "de_tabela": from_table,
            "para_tabela": to_table,
            "de_coluna": relation.get('fromColumn', ''),
            "para_coluna": relation.get('toColumn', '')
        })
    return relacionamentos


# ===================================================================
# FUNÇÃO DE ANÁLISE COM INTELIGÊNCIA ARTIFICIAL
# ===================================================================

def analisar_item_com_ia(item_dados: str, tipo_item: str, chave_api: str) -> str:
    """Envia dados de um item específico para a IA e retorna uma análise."""
    try:
        genai.configure(api_key=chave_api)
        model = genai.GenerativeModel('gemini-1.5-flash-latest')

        prompts = {
            "pagina": f"""
                Analise os dados da seguinte página de um relatório Power BI.
                Descreva, em um parágrafo curto e em linguagem de negócios, qual o provável propósito desta página,
                com base nos tipos de visuais e nas medidas utilizadas.

                DADOS DA PÁGINA:
                {item_dados}
            """,
            "tabela": f"""
                Analise a estrutura da seguinte tabela de um modelo de dados Power BI.
                Descreva, em um parágrafo curto, o propósito provável desta tabela (ex: é uma tabela de fatos, dimensão de clientes, etc.)
                e que tipo de informação ela armazena, com base em suas colunas e medidas.

                DADOS DA TABELA:
                {item_dados}
            """,
            "relacionamento": f"""
                Analise os dados a seguir sobre os relacionamentos do modelo de dados.
                Escreva uma frase explicando o que este conjunto de relacionamentos significa para o modelo de dados.

                DADOS DOS RELACIONAMENTOS:
                {item_dados}
            """
        }

        prompt_final = prompts.get(tipo_item, f"Resuma os seguintes dados: {item_dados}")
        response = model.generate_content(prompt_final)
        return response.text.strip()

    except Exception as e:
        print(f"--> Erro na IA ao analisar '{tipo_item}': {e}")
        return "Análise da IA não disponível devido a um erro."


# ===================================================================
# FUNÇÕES DE GERAÇÃO DO DOCUMENTO WORD
# ===================================================================

def salvar_versao(salvar_path):
    """Cria um nome de arquivo com versão incremental se ele já existir."""
    if not os.path.exists(salvar_path):
        return salvar_path
    base, ext = os.path.splitext(salvar_path)
    versao = 2
    while os.path.exists(f"{base}_versão_{versao:02}{ext}"):
        versao += 1
    return f"{base}_versão_{versao:02}{ext}"


def gerar_documento(cfg_data, extracoes, modelo_path, salvar_path):
    """Gera o documento Word com as descrições nos locais apropriados."""
    document = Document(modelo_path)

    for para in document.paragraphs:
        if "Data da documentação:" in para.text:
            para.text = f"Data da documentação: {datetime.now().strftime('%d/%m/%Y')}"
        elif "Nome do Relatório:" in para.text:
            para.text = f"Nome do Relatório: {cfg_data.nome_BI}"

    for titulo, conteudo_markdown in extracoes.items():
        for para in document.paragraphs:
            # Usando .lower() para uma comparação mais robusta
            if para.text.strip().lower() == titulo.lower():
                # Insere o conteúdo Markdown como um novo parágrafo logo abaixo do título
                p_novo = para.insert_paragraph_before(conteudo_markdown)
                p_novo.style = 'Normal'
                # Remove o parágrafo do título original que agora está depois
                p_para_remover = para._element
                p_para_remover.getparent().remove(p_para_remover)
                break

    caminho_final = salvar_versao(salvar_path)
    document.save(caminho_final)
    print(f"\nDocumentação gerada com sucesso em: {caminho_final}")


# ===================================================================
# FUNÇÃO PRINCIPAL (O MAESTRO)
# ===================================================================

def main():
    """Função principal que orquestra todo o processo."""
    print("--- INICIANDO AUTOMAÇÃO DE DOCUMENTAÇÃO DE POWER BI ---")

    # --- Carrega a chave de API do arquivo .env ---
    load_dotenv()
    SUA_CHAVE_API_SECRETA = os.getenv("GOOGLE_API_KEY")

    if not SUA_CHAVE_API_SECRETA:
        print("\n!!! ATENÇÃO: Chave de API não encontrada. Crie um arquivo .env com GOOGLE_API_KEY='sua_chave' !!!\n")
        return

    caminho_bi, nome_bi = cfg.caminho_BI, cfg.nome_BI
    modelo_path = os.path.join(cfg.caminho_modelo_word, cfg.nome_modelo_word)
    salvar_path = os.path.join(cfg.caminho_documentação, f'{nome_bi}_doc.docx')
    arquivo_pbit = os.path.join(caminho_bi, f'{nome_bi}.pbit')
    arquivo_zip = os.path.join(caminho_bi, f'{nome_bi}.zip')

    # --- Parte 2: Extração de Dados ---
    verificar_ou_renomear_arquivo(arquivo_pbit, arquivo_zip)
    extrair_arquivos_zip(arquivo_zip, caminho_bi, ['Report/Layout', 'DataModelSchema'])
    layout_data = carregar_dados_json(os.path.join(caminho_bi, 'Report/Layout'))
    model_data = carregar_dados_json(os.path.join(caminho_bi, 'DataModelSchema'))
    os.rename(arquivo_zip, arquivo_pbit)
    print("Dados extraídos e arquivos restaurados.")

    # --- Parte 3: Análise Item por Item com a IA ---
    print("\n--- INICIANDO ANÁLISE COM IA (ISSO PODE LEVAR ALGUNS MINUTOS) ---")

    # 1. ANÁLISE DAS PÁGINAS
    conteudo_final_paginas = []
    lista_paginas = extrair_dados_paginas(layout_data)
    for pagina in lista_paginas:
        dados_str = f"Nome da Página: {pagina['nome']}\n"
        dados_str += "Visuais:\n" + "\n".join(
            [f"  - Tipo: {v['tipo']}, Medidas: {', '.join(v['medidas'])}" for v in pagina['visuais']])

        print(f"Analisando Página: {pagina['nome']}...")
        analise_ia = analisar_item_com_ia(dados_str, "pagina", SUA_CHAVE_API_SECRETA)

        bloco_pagina = f"PÁGINA: {pagina['nome'].upper()}\n\nAnálise IA:\n{analise_ia}\n\nDetalhes Técnicos:\n{dados_str}"
        conteudo_final_paginas.append(bloco_pagina)

    # 2. ANÁLISE DAS TABELAS
    conteudo_final_tabelas = []
    lista_tabelas = extrair_dados_tabelas(model_data)
    for tabela in lista_tabelas:
        if not tabela['colunas']: continue  # Pula tabelas vazias ou de medidas

        dados_str = f"Nome da Tabela: {tabela['nome']}\n"
        dados_str += "Colunas:\n" + "\n".join(
            [f"  - {c['nome']} ({c['tipo_dado']}), Calculada: {c['calculada']}" for c in tabela['colunas']])
        if tabela['medidas']:
            dados_str += "\nMedidas na Tabela:\n" + "\n".join(
                [f"  - {m['nome']}: {m['expressao'][:100]}..." for m in tabela['medidas']])

        print(f"Analisando Tabela: {tabela['nome']}...")
        analise_ia = analisar_item_com_ia(dados_str, "tabela", SUA_CHAVE_API_SECRETA)

        bloco_tabela = f"TABELA: {tabela['nome'].upper()}\n\nAnálise IA:\n{analise_ia}\n\nDetalhes Técnicos:\n{dados_str}"
        conteudo_final_tabelas.append(bloco_tabela)

    # 3. ANÁLISE DOS RELACIONAMENTOS (um resumo do todo)
    lista_relacionamentos = extrair_dados_relacionamentos(model_data)
    if lista_relacionamentos:
        dados_str = "\n".join(
            [f"- De '{r['de_tabela']}' ({r['de_coluna']}) Para '{r['para_tabela']}' ({r['para_coluna']})" for r in
             lista_relacionamentos])
        print("Analisando Relacionamentos...")
        analise_ia = analisar_item_com_ia(dados_str, "relacionamento", SUA_CHAVE_API_SECRETA)
        conteudo_final_relacionamentos = f"Análise Geral da IA:\n{analise_ia}\n\nDetalhes Técnicos:\n{dados_str}"
    else:
        conteudo_final_relacionamentos = "Nenhum relacionamento encontrado no modelo."

    # --- Parte 4: Montando o Dicionário Final para o Word ---
    dados_finais_para_doc = {
        "Páginas": "\n\n" + ("-" * 40) + "\n\n".join(conteudo_final_paginas),
        "Tabelas": "\n\n" + ("-" * 40) + "\n\n".join(conteudo_final_tabelas),
        "Relacionamentos": conteudo_final_relacionamentos
    }

    # --- Parte 5: Geração do Documento Final ---
    print("\n--- GERANDO DOCUMENTO WORD FINAL ---")
    gerar_documento(cfg, dados_finais_para_doc, modelo_path, salvar_path)
    print("\n--- PROCESSO FINALIZADO ---")


# ===================================================================
# PONTO DE ENTRADA DO SCRIPT
# ===================================================================

if __name__ == '__main__':
    main()