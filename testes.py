"""
Automação Avançada para Geração de Documentação de Projetos Power BI
---------------------------------------------------------------------
Versão: 3.1 (Markdown → Word formatado com pypandoc)
"""

from dotenv import load_dotenv
import json
import os
import zipfile
from datetime import datetime
from docx import Document
import google.generativeai as genai
import pypandoc
import config as cfg


# ===================================================================
# FUNÇÕES UTILITÁRIAS
# ===================================================================

def verificar_ou_renomear_arquivo(arquivo_pbit, arquivo_zip):
    """Verifica se o arquivo .zip já existe; se não, renomeia o .pbit."""
    if os.path.exists(arquivo_zip):
        i = 2
        base, ext = os.path.splitext(arquivo_zip)
        while os.path.exists(f'{base}_v0{i}{ext}'):
            i += 1
        novo_nome = f'{base}_v0{i}{ext}'
        print(f"Arquivo {arquivo_pbit} renomeado para {novo_nome}.")
        os.rename(arquivo_pbit, novo_nome)
    else:
        print(f"Renomeando '{arquivo_pbit}' para '{arquivo_zip}'.")
        os.rename(arquivo_pbit, arquivo_zip)


def extrair_arquivos_zip(arquivo_zip, caminho_bi, arquivos_para_extrair):
    """Extrai arquivos específicos de um arquivo ZIP."""
    print(f"Extraindo arquivos de '{arquivo_zip}'...")
    with zipfile.ZipFile(arquivo_zip, 'r') as zip_ref:
        for arquivo in arquivos_para_extrair:
            zip_ref.extract(arquivo, caminho_bi)
    print("Arquivos extraídos com sucesso.")


def carregar_dados_json(arquivo: str, encoding: str = 'utf-16-le') -> dict:
    """Carrega dados de um arquivo JSON."""
    try:
        with open(arquivo, 'r', encoding=encoding) as f:
            return json.load(f)
    except Exception as e:
        print(f"Erro ao carregar JSON: {arquivo} - {e}")
        return {}


# ===================================================================
# EXTRAÇÃO DE DADOS
# ===================================================================

def extrair_dados(layout, model_data):
    """Extrai todas as informações do relatório Power BI de uma vez."""
    resumo = {"paginas": [], "tabelas": [], "relacionamentos": []}

    # --- PÁGINAS ---
    for section in layout.get('sections', []):
        pagina = {"nome": section.get('displayName', 'Sem Nome'), "visuais": []}
        for container in section.get("visualContainers", []):
            try:
                config_data = json.loads(container.get("config", "{}"))
                visual_type = config_data.get("singleVisual", {}).get("visualType")
                medidas = [item.get("queryRef") for items in
                           config_data.get("singleVisual", {}).get("projections", {}).values()
                           for item in items if item.get("queryRef")]
                pagina["visuais"].append({
                    "tipo": visual_type,
                    "medidas": medidas if medidas else ["Nenhuma medida explícita"]
                })
            except Exception:
                continue
        resumo["paginas"].append(pagina)

    # --- TABELAS ---
    for table in model_data.get('model', {}).get('tables', []):
        if table.get("name", "").startswith(("DateTableTemplate", "LocalDateTable")):
            continue
        tabela = {"nome": table.get("name", ""), "colunas": [], "medidas": []}
        for column in table.get('columns', []):
            tabela["colunas"].append({
                "nome": column.get("name", ""),
                "tipo_dado": column.get('dataType', ""),
                "calculada": 'Sim' if column.get('type', "") in ['calculatedTableColumn', 'calculated'] else 'Não'
            })
        for measure in table.get('measures', []):
            expr = measure.get('expression', '')
            if isinstance(expr, list):
                expr = ' '.join(filter(str.strip, expr))
            tabela["medidas"].append({
                "nome": measure.get('name', ''),
                "expressao": expr
            })
        resumo["tabelas"].append(tabela)

    # --- RELACIONAMENTOS ---
    for rel in model_data.get('model', {}).get('relationships', []):
        if rel.get('fromTable', '').startswith(("DateTableTemplate", "LocalDateTable")):
            continue
        resumo["relacionamentos"].append({
            "de_tabela": rel.get('fromTable'),
            "para_tabela": rel.get('toTable'),
            "de_coluna": rel.get('fromColumn'),
            "para_coluna": rel.get('toColumn')
        })

    return resumo


# ===================================================================
# ANÁLISE ÚNICA COM IA
# ===================================================================

def analisar_relatorio_completo(dados_relatorio, chave_api):
    """Gera uma única análise de IA com todas as informações do relatório."""
    genai.configure(api_key=chave_api)
    model = genai.GenerativeModel(
        'gemini-pro-latest',
        generation_config={
            'max_output_tokens': 8000,  # Aumenta o limite de tokens
            'temperature': 0.3,  # Mantém respostas mais técnicas/objetivas
        }
    )

    prompt = f"""
Você é um assistente técnico especializado em Power BI.
Analise o relatório completo a seguir, descrevendo:
1. O objetivo geral do relatório;
2. A função provável de cada página;
3. A função de cada tabela (ex: fato, dimensão);
4. A estrutura geral de relacionamentos;
5. Quais indicadores de negócio ou KPIs podem estar sendo apresentados.

Forneça o texto em **Markdown técnico**, com títulos e subtítulos bem estruturados.

=== DADOS COMPLETOS DO RELATÓRIO ===
{json.dumps(dados_relatorio, indent=2, ensure_ascii=False)}
"""

    print("Enviando análise global para a IA... (isso pode levar alguns segundos)")
    resposta = model.generate_content(prompt)
    return resposta.text.strip() if resposta and resposta.text else "Sem resposta da IA."


# ===================================================================
# GERAÇÃO DO DOCUMENTO
# ===================================================================

def salvar_versao(salvar_path):
    """Cria um nome de arquivo com versão incremental se ele já existir."""
    if not os.path.exists(salvar_path):
        return salvar_path
    base, ext = os.path.splitext(salvar_path)
    versao = 2
    while os.path.exists(f"{base}_versão_{versao:02}{ext}"):
        versao += 1
    return f"{base}_versão_{versao:02}{ext}"


def gerar_documento(cfg_data, texto_ia, modelo_path, salvar_path):
    """Gera o documento Word com formatação preservada (Markdown → DOCX)."""
    # Cria documento temporário com o markdown formatado
    temp_docx = os.path.join(cfg_data.caminho_documentação, "_temp_doc.docx")

    pypandoc.convert_text(
        texto_ia,
        'docx',
        format='md',
        outputfile=temp_docx,
        extra_args=['--standalone']
    )

    # Abre o modelo e anexa o conteúdo gerado
    base_doc = Document(modelo_path)
    for para in base_doc.paragraphs:
        if "Data da documentação:" in para.text:
            para.text = f"Data da documentação: {datetime.now().strftime('%d/%m/%Y')}"
        elif "Nome do Relatório:" in para.text:
            para.text = f"Nome do Relatório: {cfg_data.nome_BI}"

    # Adiciona uma quebra de página e o documento convertido
    base_doc.add_page_break()
    base_doc.add_heading("Análise Técnica Automatizada (IA)", level=1)

    temp_doc = Document(temp_docx)
    for elem in temp_doc.paragraphs:
        novo_para = base_doc.add_paragraph(elem.text)
        novo_para.style = elem.style

    caminho_final = salvar_versao(salvar_path)
    base_doc.save(caminho_final)
    os.remove(temp_docx)
    print(f"\n✅ Documentação gerada com sucesso em: {caminho_final}")


# ===================================================================
# MAIN
# ===================================================================

def main():
    print("--- INICIANDO AUTOMAÇÃO GLOBAL DE DOCUMENTAÇÃO POWER BI ---")

    load_dotenv()
    chave_api = os.getenv("GOOGLE_API_KEY")
    if not chave_api:
        print("❌ Chave da API não encontrada no .env.")
        return

    caminho_bi, nome_bi = cfg.caminho_BI, cfg.nome_BI
    modelo_path = os.path.join(cfg.caminho_modelo_word, cfg.nome_modelo_word)
    salvar_path = os.path.join(cfg.caminho_documentação, f'{nome_bi}_doc.docx')
    arquivo_pbit = os.path.join(caminho_bi, f'{nome_bi}.pbit')
    arquivo_zip = os.path.join(caminho_bi, f'{nome_bi}.zip')

    verificar_ou_renomear_arquivo(arquivo_pbit, arquivo_zip)
    extrair_arquivos_zip(arquivo_zip, caminho_bi, ['Report/Layout', 'DataModelSchema'])
    layout_data = carregar_dados_json(os.path.join(caminho_bi, 'Report/Layout'))
    model_data = carregar_dados_json(os.path.join(caminho_bi, 'DataModelSchema'))
    os.rename(arquivo_zip, arquivo_pbit)

    print("Extração concluída. Preparando dados para análise...")
    dados_relatorio = extrair_dados(layout_data, model_data)
    texto_ia = analisar_relatorio_completo(dados_relatorio, chave_api)

    print("\nGerando documento final...")
    gerar_documento(cfg, texto_ia, modelo_path, salvar_path)

    print("\n--- ✅ PROCESSO FINALIZADO COM SUCESSO ---")


if __name__ == '__main__':
    main()
